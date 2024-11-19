import os
import time
import subprocess
from collections import defaultdict
from pynput import mouse, keyboard
import threading
from PIL import Image
from docx import Document
from docx.shared import Inches
import pyautogui
import cv2
import numpy as np
import win32gui
import xml.etree.ElementTree as ET
import tkinter as tk

# Constants
#SCREENSHOT_DIR = "screenshots"
#TEMPLATE_PATH = "C:\\Users\\cogni\\Downloads\\13_11_2024 at 13_54.docx"
#OUTPUT_PDD_PATH = "final_process_definition_document.docx"
#DRAWIO_CLI_PATH = "C:/Program Files/draw.io/draw.io.exe"
#FLOWCHART_PATH = "flowchart.drawio""""


# Directories and Paths
screenshot_dir = "screenshots"
template_path = "C:\\Users\\cogni\\Downloads\\13_11_2024 at 13_54.docx"  # Template file path
pdd_output_path = 'final_process_definition_document.docx'  # Output PDD file
flowchart_path = "flowchart.drawio"
DRAWIO_CLI_PATH = "C:/Program Files/draw.io/draw.io.exe"
if not os.path.exists(screenshot_dir):
    os.makedirs(screenshot_dir)

# Global Variables
actions = []
screenshot_counter = 0
unique_screenshots = []
window_stats = defaultdict(lambda: {"mouse_clicks": 0, "text_entries": 0, "keys_pressed": 0, "hotkeys_used": 0})
hotkeys_used = 0
pressed_keys = set()
text_buffer = ""
is_text_entry_active = False
action_start_time = None
recording = False

# Capture Screenshot
def capture_screenshot():
    screenshot = pyautogui.screenshot()
    screenshot_cv = np.array(screenshot)
    screenshot_cv = cv2.cvtColor(screenshot_cv, cv2.COLOR_RGB2BGR)
    return screenshot_cv

# Annotate Screenshot
def annotate_screenshot(image, coordinates):
    x1, y1 = coordinates[0]
    x2, y2 = coordinates[1]
    annotated_image = cv2.rectangle(image, (x1, y1), (x2, y2), (0, 255, 0), 2)
    return annotated_image

# Save Annotated Image
def save_annotated_image(image, filename):
    annotated_image_pil = Image.fromarray(cv2.cvtColor(image, cv2.COLOR_BGR2RGB))
    annotated_image_pil.save(filename)

# Get Active Window Title
def get_active_window_title():
    window = win32gui.GetForegroundWindow()
    return win32gui.GetWindowText(window)

# Update Window Statistics
def update_window_statistics(window_name, action_type):
    window_stats[window_name][action_type] += 1

# Mouse Click Listener
def on_click(x, y, button, pressed):
    global actions, screenshot_counter, action_start_time
    if pressed and recording:
        action_end_time = time.time()
        elapsed_time = action_end_time - action_start_time if action_start_time else 0
        action_start_time = action_end_time
        action = "Mouse clicked"
        active_window = get_active_window_title()
        
        # Capture Screenshot
        screenshot = capture_screenshot()
        annotated_image = annotate_screenshot(screenshot, [(x - 50, y - 50), (x + 50, y + 50)])
        screenshot_counter += 1
        filename = os.path.join(screenshot_dir, f"annotated_screenshot_{screenshot_counter}.png")
        save_annotated_image(annotated_image, filename)
        
        unique_screenshots.append({"filename": filename, "action": action, "window": active_window, "time": elapsed_time})
        actions.append({"action": action, "window": active_window, "time": elapsed_time})
        update_window_statistics(active_window, "mouse_clicks")

# Keyboard Listener for Consolidated Text Entry
def on_press(key):
    global action_start_time, hotkeys_used, pressed_keys, text_buffer, actions, is_text_entry_active
    if recording:
        action_end_time = time.time()
        elapsed_time = action_end_time - action_start_time if action_start_time else 0
        action_start_time = action_end_time
        active_window = get_active_window_title()

        # Handle key inputs for text entries
        try:
            # Check if it's a regular character key
            if hasattr(key, 'char') and key.char is not None:
                text_buffer += key.char  # Add character to buffer
                is_text_entry_active = True
            # Handle text commit keys like Enter, Tab, or Space
            elif key in {keyboard.Key.enter, keyboard.Key.tab, keyboard.Key.space} and text_buffer:
                action = f"Text entered"
                actions.append({"action": action, "window": active_window, "time": elapsed_time})
                update_window_statistics(active_window, "text_entries")
                text_buffer = ""  # Reset buffer after committing the text entry
                is_text_entry_active = False
        except AttributeError:
            # Handle special keys, potentially part of hotkeys
            pressed_keys.add(str(key))
            if is_text_entry_active:  # Commit any buffer contents if a special key is pressed
                action = f"Text entered"
                actions.append({"action": action, "window": active_window, "time": elapsed_time})
                update_window_statistics(active_window, "text_entries")
                text_buffer = ""
                is_text_entry_active = False

        # Check for hotkeys (Ctrl, Alt, Shift combinations)
        if any(k in pressed_keys for k in ["Key.ctrl", "Key.alt", "Key.shift"]):
            action = f"Hotkey {' + '.join(pressed_keys)}"
            hotkeys_used += 1
            actions.append({"action": action, "window": active_window, "time": elapsed_time})
            update_window_statistics(active_window, "hotkeys_used")
        else:
            # For individual key presses not part of a text entry
            action = f"Key {key} pressed"
            actions.append({"action": action, "window": active_window, "time": elapsed_time})
            update_window_statistics(active_window, "keys_pressed")

# Reset pressed_keys on key release
def on_release(key):
    try:
        if hasattr(key, 'char') and key.char:
            pressed_keys.discard(key.char)
        else:
            pressed_keys.discard(str(key))
    except KeyError:
        pass


# Load Template Document
def load_template():
    return Document(template_path)


# Generate Statistics and Fill into Template
def generate_statistics_in_template(doc):
    # Locate the tables in the template to insert statistics
    high_level_table = doc.tables[3]  # Assuming first table is for high-level stats
    detailed_table = doc.tables[4]    # Assuming second table is for detailed stats

    # High-Level Statistics
    high_level_data = [
        ["Processes", "Windows", "Actions", "Mouse clicks", "Keys pressed", "Text entries", "Hotkeys used", "Time"],
        [1, len(window_stats), len(actions), sum(w["mouse_clicks"] for w in window_stats.values()),
         sum(w["keys_pressed"] for w in window_stats.values()), sum(w["text_entries"] for w in window_stats.values()),
         hotkeys_used, f"{sum(a['time'] for a in actions):.1f} sec"]
    ]

    # Populate High-Level Stats
    for i, row in enumerate(high_level_data):
        for j, cell_text in enumerate(row):
            # Add a new row if beyond current row count
            if i >= len(high_level_table.rows):
                high_level_table.add_row()
            high_level_table.cell(i, j).text = str(cell_text)

    # Detailed Statistics
    detailed_data = [
        ["Window Name", "Mouse Clicks", "Text Entries", "Keys Pressed"],
        *[[window, data["mouse_clicks"], data["text_entries"], data["keys_pressed"]] for window, data in window_stats.items()]
    ]

    # Populate Detailed Stats
    for i, row in enumerate(detailed_data):
        for j, cell_text in enumerate(row):
            # Add a new row if beyond current row count
            if i >= len(detailed_table.rows):
                detailed_table.add_row()
            detailed_table.cell(i, j).text = str(cell_text)

# Flowchart Generation

# Helper: Generate Flowchart in .drawio Format
# Helper: Generate Flowchart in .drawio Format
def generate_drawio_flowchart(actions, file_path):
    print("Generating flowchart in .drawio format...")
    mxfile = ET.Element("mxfile")
    diagram = ET.SubElement(mxfile, "diagram", name="High Level Process")
    mxGraphModel = ET.SubElement(diagram, "mxGraphModel")
    root = ET.SubElement(mxGraphModel, "root")

    # Base structure
    ET.SubElement(root, "mxCell", id="0")
    ET.SubElement(root, "mxCell", id="1", parent="0")

    x, y = 200, 50
    step_width, step_height = 300, 100
    spacing = 50
    previous_id = "1"

    # Group actions by window and summarize them
    grouped_actions = defaultdict(lambda: {"count": 0, "time": 0})
    for action in actions:
        grouped_actions[action["window"]]["count"] += 1
        grouped_actions[action["window"]]["time"] += action["time"]

    for i, (window, summary) in enumerate(grouped_actions.items()):
        box_id = str(i + 2)
        action_label = (
            f"Window: {window}\n"
            f"Actions: {summary['count']}\n"
            f"Est. Time: {summary['time']:.2f}s"
        )
        
        # Add flowchart box
        box = ET.SubElement(root, "mxCell", id=box_id, value=action_label, style="rounded=1;whiteSpace=wrap;html=1;",
                            vertex="1", parent="1")
        ET.SubElement(box, "mxGeometry", x=str(x), y=str(y), width=str(step_width), height=str(step_height), 
                      **{"as": "geometry"})
        
        # Add connector
        ET.SubElement(root, "mxCell", id=f"{previous_id}-{box_id}", edge="1", source=previous_id, 
                      target=box_id, parent="1")
        
        y += step_height + spacing
        previous_id = box_id

    tree = ET.ElementTree(mxfile)
    tree.write(file_path, encoding="utf-8", xml_declaration=True)
    print(f"Flowchart saved as .drawio at {file_path}")
# Helper: Export Flowchart to PNG
def export_flowchart_to_png():
    print("Exporting flowchart to PNG...")
    input_path = flowchart_path
    output_path = flowchart_path.replace(".drawio", ".png")

    if not os.path.exists(DRAWIO_CLI_PATH):
        raise FileNotFoundError("draw.io CLI not found. Please verify the DRAWIO_CLI_PATH.")
    if not os.path.exists(input_path):
        raise FileNotFoundError("The .drawio file does not exist. Flowchart generation might have failed.")

    # Run draw.io export command
    subprocess.run([DRAWIO_CLI_PATH, "-x", "-f", "png", input_path, "-o", output_path], check=True)
    print(f"Flowchart exported to PNG at {output_path}")




# Prompt user to edit the flowchart
def prompt_for_flowchart_edit():
    print(f"Flowchart generated at {flowchart_path}. Please edit it as needed.")
    input("Press Enter to continue after editing the flowchart...")

# Generate PDD
def generate_pdd():
    doc = load_template()

    # Insert flowchart into the PDD
    prompt_for_flowchart_edit()  # Prompt user to edit flowchart before proceeding
    export_flowchart_to_png()  # Fit flowchart to specific size

    # Rest of the PDD generation process
    # Insert screenshots and actions in a table format
    screenshot_table = doc.tables[5]  # Assuming a specific location in the template
    for screenshot in unique_screenshots:
        row = screenshot_table.add_row().cells
        row[0].text = screenshot['window']
        row[1].text = screenshot['action']
        row[2].text = f"{screenshot['time']:.1f} sec"
        row[3].text = ""  # Placeholder for screenshot image
        row[3].paragraphs[0].add_run().add_picture(screenshot['filename'], width=Inches(2.5))

    
    # Populate statistics in the PDD template
    generate_statistics_in_template(doc)

    # Insert flowchart into the "High Level Process Map" section
    for paragraph in doc.paragraphs:
        if "High Level Process Map" in paragraph.text:
            flowchart_image_path = flowchart_path.replace(".drawio", ".png")
            paragraph.add_run("\n").add_picture(flowchart_image_path, width=Inches(6.5))
            print("Flowchart inserted into PDD.")
            break

    # Save the document
    doc.save(pdd_output_path)
    print("PDD generated successfully.")


# GUI and Main Functions
def start_recording():
    global recording, action_start_time
    recording = True
    action_start_time = time.time()

# Stop recording and handle flowchart edits
def stop_recording():
    global recording
    recording = False
    generate_drawio_flowchart(actions, flowchart_path)
    print("Flowchart generation complete. Please edit the flowchart if needed.")
    generate_pdd()

def create_gui():
    root = tk.Tk()
    root.title("Process Recorder")
    tk.Button(root, text="Start Recording", command=start_recording, bg="green", fg="white", width=20).pack(pady=10)
    tk.Button(root, text="Stop Recording", command=stop_recording, bg="red", fg="white", width=20).pack(pady=10)
    root.mainloop()

def listener_threads():
    mouse.Listener(on_click=on_click).start()
    keyboard.Listener(on_press=on_press).start()

# Main function
def main():
    threading.Thread(target=listener_threads, daemon=True).start()
    create_gui()

if __name__ == "__main__":
    main()